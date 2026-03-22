"""
Módulo de Escritura en Word para el Reporte Estándar.
Inserta tabla + gráfico (y texto opcional) en anclas del documento.
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.table import Table


def find_paragraph(doc: Document, text_to_find: str) -> Paragraph | None:
    text_to_find_lower = text_to_find.lower()
    for p in doc.paragraphs:
        if text_to_find_lower in p.text.lower():
            return p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if text_to_find_lower in p.text.lower():
                        return p
    return None


def _apply_table_style_safely(tbl: Table):
    preferred_styles = [
        "Table Grid",
        "Grid Table 1 Light",
        "Light Grid",
        "Cuadrícula clara",
        "Tabla con cuadrícula",
    ]
    for style_name in preferred_styles:
        try:
            tbl.style = style_name
            return
        except KeyError:
            continue
        except Exception:
            continue
    print("[yellow]Advertencia: Ninguno de los estilos de tabla preferidos se encontró.[/yellow]")


def _to_table_df(data) -> pd.DataFrame | None:
    """
    Convierte Series/DataFrame/dict(simple) a DataFrame para tabla.
    """
    if isinstance(data, pd.Series):
        df = data.reset_index()
        df.columns = ["Categoría", "Cantidad"]
        return df

    if isinstance(data, pd.DataFrame):
        return data.reset_index()

    # soporte básico para dict {k:v}
    if isinstance(data, dict):
        try:
            s = pd.Series(data)
            df = s.reset_index()
            df.columns = ["Categoría", "Cantidad"]
            return df
        except Exception:
            return None

    return None


def _insert_table(doc: Document, anchor_paragraph: Paragraph, data):
    """
    Inserta tabla antes del párrafo ancla.
    """
    df = _to_table_df(data)
    if df is None:
        print("[red]Error: El tipo de dato para la tabla no es válido.[/red]")
        return

    num_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=num_cols)
    _apply_table_style_safely(table)

    # Header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)

    # Rows
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = "" if pd.isna(val) else str(val)

    # mover tabla antes del ancla (lxml)
    anchor_paragraph._p.addprevious(table._tbl)


def _insert_image(anchor_paragraph: Paragraph, image_path: Path):
    """
    Inserta imagen antes del párrafo ancla.
    """
    p = anchor_paragraph.insert_paragraph_before()
    p.alignment = 1  # center
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))
    p.insert_paragraph_before()  # espacio


def insert_content_at_anchor(
    doc: Document,
    anchor_text: str,
    data,
    image_path: Path,
    summary_text: str | None = None,
    *,
    remove_anchor: bool = True,
):
    """
    Inserta un bloque (tabla + imagen + texto opcional) sobre un ancla.
    """
    anchor_paragraph = find_paragraph(doc, anchor_text)
    if not anchor_paragraph:
        print(f"[bold yellow]Advertencia: No se encontró el ancla '{anchor_text}'.[/bold yellow]")
        return

    if summary_text:
        p_text = anchor_paragraph.insert_paragraph_before(summary_text)
        p_text.insert_paragraph_before()

    _insert_image(anchor_paragraph, image_path)
    _insert_table(doc, anchor_paragraph, data)

    if remove_anchor:
        p = anchor_paragraph._element
        p.getparent().remove(p)


def insert_multiple_contents_at_anchor(
    doc: Document,
    anchor_text: str,
    items: list[dict],
    *,
    remove_anchor: bool = True,
):
    """
    Inserta múltiples bloques en una sola ancla.
    Cada item: {"data": ..., "image_path": Path, "summary_text": str|None}
    """
    anchor_paragraph = find_paragraph(doc, anchor_text)
    if not anchor_paragraph:
        print(f"[bold yellow]Advertencia: No se encontró el ancla '{anchor_text}'.[/bold yellow]")
        return

    # Reversa para conservar orden visual en inserción-before
    for item in reversed(items):
        data = item.get("data")
        image_path = item.get("image_path")
        summary_text = item.get("summary_text")

        if summary_text:
            p_text = anchor_paragraph.insert_paragraph_before(summary_text)
            p_text.insert_paragraph_before()

        _insert_image(anchor_paragraph, image_path)
        _insert_table(doc, anchor_paragraph, data)

    if remove_anchor:
        p = anchor_paragraph._element
        p.getparent().remove(p)