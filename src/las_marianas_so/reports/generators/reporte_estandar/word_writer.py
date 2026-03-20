"""
Utilidades de bajo nivel para manipular documentos Word.
Versión completa y fiel adaptada del repositorio 'temporal'.
"""
import re
import unicodedata
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

def open_template(template_path: Path) -> Document:
    """Abre un documento de plantilla desde una ruta."""
    if not template_path.exists():
        raise FileNotFoundError(f"No existe la plantilla en la ruta: {template_path}")
    return Document(str(template_path))

def save(doc: Document, report_path: Path) -> None:
    """Guarda un documento en la ruta de salida, creando directorios si es necesario."""
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(report_path))

def _norm_text(s: str) -> str:
    """Normaliza texto para búsqueda (ignora acentos, mayúsculas, espacios extra)."""
    s = str(s or "").replace("\u00a0", " ").strip().casefold()
    s = "".join(ch for ch in unicodedata.normalize("NFD", s) if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", s)

def _iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Itera sobre todos los párrafos del documento, incluyendo los de las tablas."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs

def find_paragraph_smart(doc: Document, needle: str, start_after_text: str | None = None) -> Optional[Paragraph]:
    """
    Busca un párrafo que contenga un texto ('needle').
    Opcionalmente, puede empezar la búsqueda después de otro texto ('start_after_text').
    """
    needle_norm = _norm_text(needle)
    start_norm = _norm_text(start_after_text) if start_after_text else None
    
    started = not bool(start_norm)
    for p in _iter_all_paragraphs(doc):
        p_text_norm = _norm_text(p.text)
        if not started and start_norm in p_text_norm:
            started = True
            continue
        if started and needle_norm in p_text_norm:
            return p
    return None

def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    """Inserta un párrafo vacío después de un párrafo ancla."""
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)

def insert_paragraph_after_paragraph(anchor: Paragraph, text: str, *, bold: bool = False) -> Paragraph:
    """Inserta un párrafo con texto después de un párrafo ancla."""
    p = _insert_paragraph_after(anchor)
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def insert_picture_after_paragraph(anchor: Paragraph, image_path: Path, width_inches: float) -> Paragraph:
    """Inserta una imagen centrada después de un párrafo ancla."""
    p = _insert_paragraph_after(anchor)
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    # Inserta un párrafo vacío después de la imagen para espaciar
    return _insert_paragraph_after(p)

def insert_table_after_paragraph(doc: Document, anchor: Paragraph, headers: list, rows: list) -> Table:
    """Inserta una tabla con estilo después de un párrafo ancla."""
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    
    # Rellenar tabla
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data or "")

    # Mover la tabla justo después del párrafo ancla
    anchor._p.addnext(table._tbl)
    _insert_paragraph_after(anchor.next_sibling) # Párrafo de espacio
    return table