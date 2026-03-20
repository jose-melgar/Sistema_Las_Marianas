"""
Core del Sistema de Generación de Reportes (Actúa como Router).
"""
from pathlib import Path
from docx import Document
from rich.console import Console

# --- NUEVA IMPORTACIÓN DEL MANEJADOR ---
from .handlers import standard_report_handler

console = Console()

def _find_text_in_doc(doc: Document, text_to_find: str) -> bool:
    """Busca un texto específico dentro de todo el documento Word."""
    text_to_find_lower = text_to_find.lower()
    for p in doc.paragraphs:
        if text_to_find_lower in p.text.lower(): return True
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if text_to_find_lower in p.text.lower(): return True
    return False

# --- El nombre de la función se mantiene igual para no romper la UI ---
def generate_standard_report(obra: str, year: int, month: int, report_type: str):
    """
    Punto de entrada que valida y enruta al manejador de reporte apropiado.
    """
    base_path = Path(__file__).resolve().parent.parent.parent.parent
    template_path = base_path / "templates/informes/prueba.docx"
    output_path_dir = base_path / "outputs"
    output_path_dir.mkdir(exist_ok=True)
    
    if not template_path.exists():
        console.print(f"[bold red]Error: No se encuentra la plantilla en:[/bold red] {template_path}")
        return

    doc = Document(template_path)
    
    if not _find_text_in_doc(doc, report_type):
        console.print(f"[bold red]Error de Plantilla:[/bold red] El archivo no es válido para un '{report_type}'.")
        return
        
    console.print(f"[bold green]Validación de plantilla exitosa.[/bold green]")
    
    # --- LÓGICA DE ENRUTAMIENTO ---
    # Decide a qué manejador llamar basado en el 'report_type'.
    if report_type == "Reporte Estandar":
        # Llama al manejador del reporte estándar y le pasa el documento
        standard_report_handler.run(doc)
    
    elif report_type == "Reporte de Vulnerabilidad":
        console.print("[yellow]El 'Reporte de Vulnerabilidad' aún no está implementado.[/yellow]")
        # Aquí se llamaría a 'vulnerability_report_handler.run(doc)' en el futuro.
    else:
        console.print(f"[bold red]Error: Tipo de reporte '{report_type}' desconocido.[/bold red]")
        return
    
    # --- GUARDADO FINAL ---
    # Después de que el manejador ha modificado el 'doc', lo guardamos.
    output_filename = f"Reporte_{obra.replace(' ', '_')}_{year}_{month:02d}.docx"
    output_path = output_path_dir / output_filename
    doc.save(output_path)
    
    console.print(f"\n[bold green]¡Reporte generado![/bold green] Puedes encontrarlo en: [cyan]{output_path}[/cyan]")