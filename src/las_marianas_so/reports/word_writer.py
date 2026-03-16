"""
Utilidades para la Escritura de Documentos Word con python-docx.

Adaptado del proyecto `temporal`, este módulo encapsula operaciones comunes
para crear y formatear informes en formato .docx.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_styled_document() -> Document:
    """Crea un documento con estilos básicos predefinidos."""
    doc = Document()
    # Estilos (se pueden ampliar)
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(11)
    return doc

def add_main_title(doc: Document, text: str):
    """Añade un título principal centrado."""
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)

def add_section_title(doc: Document, text: str):
    """Añade un título de sección."""
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    p.paragraph_format.space_after = Pt(12)

def add_table_from_df(doc: Document, df, style='Table Grid'):
    """Añade una tabla a partir de un DataFrame de pandas."""
    # Añadir una fila para los encabezados y luego las filas de datos
    table = doc.add_table(rows=1, cols=len(df.columns), style=style)
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Filas de datos
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    doc.add_paragraph() # Espacio después de la tabla

def save_document(doc: Document, output_path: Path, filename: str):
    """Guarda el documento Word en la ruta especificada."""
    output_path.mkdir(parents=True, exist_ok=True)
    filepath = output_path / filename
    doc.save(filepath)
    print(f"   -> Documento guardado en: {filepath}")